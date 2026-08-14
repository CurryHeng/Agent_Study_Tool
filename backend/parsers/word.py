"""Word（.docx）解析器，按标题样式切分章节。"""
from docx import Document as DocxDocument

from parsers.base import ParsedDocument, Section


def _heading_level(style_name: str) -> int:
    digits = "".join(ch for ch in style_name if ch.isdigit())
    return int(digits) if digits else 1


class WordParser:
    source_type = "word"

    def parse(self, file_path: str) -> ParsedDocument:
        doc = DocxDocument(file_path)
        sections: list[Section] = []
        current: Section | None = None

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style_name = para.style.name if para.style else ""
            if style_name.startswith(("Heading", "标题")):
                level = _heading_level(style_name)
                current = Section(title=text, level=level)
                sections.append(current)
            elif current is not None:
                current.paragraphs.append(text)

        return ParsedDocument(title="", source_type=self.source_type, sections=sections)
